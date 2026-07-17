#!/usr/bin/env python3
"""gen_heavy_docx.py -- calibrated heavy-docx generator for Phase 28 (D-07).

Generates a multi-page/table .docx into a caller-chosen path, sized by a
`--page-units` calibration knob, targeting ~200s of LibreOffice conversion
time on the in-cluster document-worker (calibrated by a live trial run --
there is no local soffice/libreoffice binary on this host, so DO NOT attempt
to time this locally; see 28-RESEARCH.md Pitfall 4).

Invoked ONLY via ephemeral uv, e.g.:
    uv run --with python-docx python3 scripts/fixtures/gen_heavy_docx.py \
        --page-units 300 --out /tmp/heavy.docx

This script imports only `docx` (python-docx) plus stdlib. It is NEVER added
to any persisted dependency manifest (no requirements.txt, no go.mod touch)
-- python-docx is pulled fresh by `uv run --with` on every invocation.

Per 28-RESEARCH.md Security Domain: this generator script is committed, but
the .docx files it produces are NOT -- they are generated to /tmp (or any
caller-chosen scratch path) at run time by the load-proof gate, never
committed to the repo.
"""
import argparse
import os
import sys

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Reuse the existing image fixture so no new binary asset needs to be
# committed for this generator (28-RESEARCH.md Code Examples).
SAMPLE_IMAGE = os.path.join("internal", "e2e", "testdata", "sample.png")


def add_toc_field(document):
    """Insert a TOC field via low-level OXML so LibreOffice does real field
    layout work when it opens/converts the document (not just static text)."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def build_document(page_units, out_path):
    d = docx.Document()

    d.add_heading("OctoConv Phase 28 Load-Proof Heavy Fixture", level=0)
    add_toc_field(d)
    d.add_page_break()

    have_sample_image = os.path.isfile(SAMPLE_IMAGE)

    for i in range(page_units):
        d.add_heading(f"Section {i}", level=1)
        d.add_paragraph("Lorem ipsum " * 80)
        table = d.add_table(rows=8, cols=6)
        for row in table.rows:
            for cell in row.cells:
                cell.text = "data " * 4
        if have_sample_image and i % 10 == 0:
            d.add_picture(SAMPLE_IMAGE)
        d.add_page_break()

    d.save(out_path)


def main():
    parser = argparse.ArgumentParser(
        description="Generate a calibrated heavy .docx fixture for the Phase 28 load-proof gate."
    )
    parser.add_argument(
        "--page-units",
        type=int,
        default=300,
        help="Number of heading+paragraph+table (+ every-10th image) units to generate (default: 300). "
        "Calibration knob -- larger N means a heavier document / longer LibreOffice conversion time.",
    )
    parser.add_argument(
        "--out",
        type=str,
        required=True,
        help="Output path for the generated .docx (never commit this file -- generate to a scratch path).",
    )
    args = parser.parse_args()

    if args.page_units < 1:
        print("FAIL: --page-units must be >= 1", file=sys.stderr)
        sys.exit(1)

    build_document(args.page_units, args.out)

    print(f"Generated heavy docx: {args.out}")
    print(f"page-units: {args.page_units}")


if __name__ == "__main__":
    main()
